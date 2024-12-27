from flask import Flask, jsonify, request, send_file
from flask_sqlalchemy import SQLAlchemy
from flask_cors import CORS
from jinja2 import Template
from docx import Document
import io
from datetime import datetime

app = Flask(__name__)
CORS(app) 
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///tasks.db'
db = SQLAlchemy(app)

class Task(db.Model):
    id = db.Column(db.Integer, primary_key = True)
    title = db.Column(db.String(100), nullable = False)
    completed = db.Column(db.Boolean, default = False)

with app.app_context():
    db.create_all()

@app.route('/tasks', methods=['GET'])
def get_tasks():
    tasks = Task.query.all()
    return jsonify([{'id': task.id, 'title': task.title, 'completed': task.completed} for task in tasks])

@app.route('/tasks', methods=['POST'])
def add_task():
    data = request.json
    new_task = Task(title = data['title'])
    db.session.add(new_task)
    db.session.commit()
    return jsonify({'id': new_task.id, 'title': new_task.title, 'completed':new_task.completed}), 201

@app.route('/tasks/<int:id>', methods=['GET', 'PUT'])
def update_task(id):
    task = Task.query.get_or_404(id)
    data = request.get_json()
    task.title = data.get('title', task.title)
    task.completed = data.get('completed', task.completed)
    db.session.commit()
    return jsonify({'id': task.id, 'title': task.title, 'completed': task.completed})

@app.route('/tasks/<int:id>', methods=['DELETE'])
def delete_task(id):
    task = Task.query.get_or_404(id)
    db.session.delete(task)
    db.session.commit()
    return jsonify({'message': 'Task deleted'})

@app.route('/download', methods=['GET'])
def download_docx():
    tasks = Task.query.all()

    doc = Document()
    doc.add_heading('To-Do list', level=1)

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Task Title'
    hdr_cells[1].text = 'Status'

    for task in tasks:
        row_cells = table.add_row().cells
        row_cells[0].text = task.title
        row_cells[1].text = 'Completed' if task.completed else 'Pending'

    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)

    return send_file(
        doc_stream,
        as_attachment=True,
        download_name=f'to-do_list_{datetime.now().strftime("%Y_%m_%d")}.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

if __name__ == '__main__':
    app.run(debug=True)



